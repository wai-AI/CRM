import tkinter as tk
from tkinter import ttk
import sqlite3
import os
from docx import Document
import datetime


def save_to_database():
    text1 = text_entry1.get() # Получаем текст из текстового поля
    text2 = text_entry2.get()
    text3 = text_entry3.get()

    selected_table = table_combobox.get() # Получаем выбранную таблицу из выпадающего списка
    # Устанавливаем соединение с базой данных
    connection = sqlite3.connect('data.db')
    cursor = connection.cursor()
    # Создаем таблицы, если они еще не существуют
    # Вставляем введенный текст в выбранную таблицу
    cursor.execute(f'''INSERT INTO {selected_table} (column1, column2, column3) VALUES (?, ?, ?)''', (text1, text2, text3,))
    connection.commit() # Фиксируем изменения
    connection.close() # Закрываем соединение
    # Очищаем текстовое поле
    text_entry1.delete(0, tk.END)
    text_entry2.delete(0, tk.END)
    text_entry3.delete(0, tk.END)

def load():
    # Устанавливаем соединение с базой данных
    connection = sqlite3.connect('data.db')
    cursor = connection.cursor()
    date = datetime.datetime.now().strftime('%Y-%m-%d')

    # Получаем данные из таблицы "secondary"
    cursor.execute('''SELECT column1 FROM secondary''')
    companies = cursor.fetchall()
    text_values = ', '.join([row[0] for row in companies])
    cursor.execute('''SELECT column2 FROM secondary''')
    contact_name = cursor.fetchall()
    name_values = ', '.join([row[0] for row in contact_name])
    cursor.execute('''SELECT column3 FROM secondary''')
    contact_num = cursor.fetchall()
    contact = ', '.join([row[0] for row in contact_num])
    connection.close()
    
    # Подсчитываем количество текстовых значений в каждом столбце


    # Записываем количество текстовых значений в каждом столбце в отдельный документ Word
    document = Document()
    document.add_heading('Количество текстовых значений в каждом столбце', level=1)
    document.add_paragraph(f'Назви компаній: {text_values}')
    document.add_paragraph(f'Контактна особа: {name_values}')
    document.add_paragraph(f'Контакт: {contact}')
    
    # Сохраняем документ
    filename = f"{date}.docx"
    if os.path.exists(filename):
        os.remove(filename)  # Удаляем файл, если он уже существует
    document.save(filename)


# Создаем основное окно
root = tk.Tk()
root.title("Перспектива CRM")

# Создаем текстовое поле для ввода текста
text_entry1 = tk.Entry(root, width=50)
label1 = tk.Label(root, text="Введіть назву компанії:")
label1.pack()
text_entry1.pack(pady=10)
text_entry2 = tk.Entry(root, width=50)
label2 = tk.Label(root, text="Введіть ПІБ контактної особи:")
label2.pack()
text_entry2.pack(pady=10)
text_entry3 = tk.Entry(root, width=50)
label3 = tk.Label(root, text="Введіть контактний номер:")
label3.pack()
text_entry3.pack(pady=10)

# Создаем выпадающий список для выбора таблицы
table_combobox = ttk.Combobox(root, values=["texts", "main", "secondary"],state="readonly")
table_combobox.current(2) # Устанавливаем значения по умолчанию
table_combobox.pack(pady=5)

# Создаем кнопку для сохранения текста в выбранную таблицу
save_button = tk.Button(root, text="Зберегти на сервері", command=save_to_database)
load_button = tk.Button(root, text="Вивантажити звіт за день", command=load)
save_button.pack(pady=5)
load_button.pack(pady=5)

# Запускаем цикл обработки событий
root.mainloop()
